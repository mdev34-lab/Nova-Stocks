import yfinance as yf
import matplotlib.pyplot as plt
import mplfinance as mpf
import pandas as pd
from PySide6.QtWidgets import QMessageBox
import docx
import datetime
from collections import Counter

HUMAN_READABLE_PERIODS = {
    "1d": "1 dia", 
    "5d": "5 dias", 
    "1mo": "1 mês", 
    "3mo": "3 meses", 
    "6mo": "6 meses", 
    "1y": "1 ano", 
    "2y": "2 anos", 
    "5y": "5 anos", 
    "max": "máximo"
}

def is_valid_ticker(symbol: str):
    if not symbol or not symbol.isalnum():
        return False
    if len(symbol) > 5 and not symbol.endswith('11'):
        return False
    try:
        symbol: str = symbol.upper()
        if not symbol.endswith('.SA'):
            symbol += '.SA'
        ticker = yf.Ticker(symbol)
        return ticker.info.get('symbol') == symbol
    except Exception as e:
        # print(f"Error validating ticker: {e}")
        return False

def _fetch(symbol, period=None, start_date=None, end_date=None, show_volume=True):
    symbol = symbol.upper()
    if not symbol.endswith('.SA'):
        symbol += '.SA'

    try:
        stock = yf.Ticker(symbol)
        if start_date and end_date:
            data = stock.history(start=start_date, end=end_date)
        else:
            data = stock.history(period=period)

        if data.empty:
            raise ValueError("No data available for this stock symbol")
        return data
    except Exception as e:
        raise ValueError(f"Error fetching data: {str(e)}")

def validate_dates(start_date, end_date):
    """
    Validate the start and end dates.
    
    Parameters
    ----------
    start_date : str
        Start date in YYYY-MM-DD format
    end_date : str
        End date in YYYY-MM-DD format
    
    Returns
    -------
    tuple
        (start_date, end_date) as datetime objects if valid
    
    Raises
    ------
    ValueError
        If dates are invalid
    """
    try:
        start = datetime.datetime.strptime(start_date, "%Y-%m-%d")
        end = datetime.datetime.strptime(end_date, "%Y-%m-%d")
        
        if start >= end:
            raise ValueError("Start date must be before end date")
            
        if end > datetime.datetime.now():
            raise ValueError("End date cannot be in the future")
            
        return start, end
    except ValueError as e:
        if "strptime" in str(e):
            raise ValueError("Invalid date format. Use YYYY-MM-DD")
        raise e

def fetch_data(symbol, period=None, start_date=None, end_date=None, show_ma_sma=False, show_ma_ema=False, show_volume=True, ma_period=20):
    try:
        data = _fetch(symbol, period=period, start_date=start_date, end_date=end_date, show_volume=show_volume)

        # Configuração do estilo do gráfico
        mc = mpf.make_marketcolors(up='g', down='r',
                                 edge='inherit',
                                 wick='inherit',
                                 volume='in')
        s = mpf.make_mpf_style(marketcolors=mc, gridstyle=':', y_on_right=False)

        # Lista para armazenar os plots adicionais
        addplots = []

        # Adiciona média móvel se solicitado
        if show_ma_sma:
            # Calcula a média móvel simples
            sma = data['Close'].rolling(window=ma_period).mean()
            addplots.append(mpf.make_addplot(sma, color='blue', width=0.8, label='SMA'))

        if show_ma_ema:
            # Calcula a média móvel exponencial
            ema = data['Close'].ewm(span=ma_period, adjust=False).mean()
            addplots.append(mpf.make_addplot(ema, color='green', width=0.8, label='EMA'))

        # Plotando o gráfico
        fig, _ = mpf.plot(data, type='candle', 
                        style=s,
                        title=f'{symbol} Preço da Ação',
                        volume=show_volume,
                        addplot=addplots if addplots else [],  # Ensure it's an empty list if no plots
                        returnfig=True)

        # Show chart in a native MPL window
        plt.show()

    except ValueError as e:
        QMessageBox.critical("Erro", str(e))

def save_data_excel(symbol, start_date, end_date, file_name, parent):
    """
    Save the stock data as an Excel spreadsheet

    Parameters
    ----------
    symbol : str
        Stock symbol
    period : str
        Data period
    file_name : str
        File name to be saved
    """
    try:
        data = _fetch(symbol, start_date=start_date, end_date=end_date, show_volume=True)

        # Translate column titles
        data.columns = ['Abertura', 'Alta', 'Baixa', 'Fechamento', 'Volume', 'Dividendos', 'Splits']

        # Create ExcelWriter object
        with pd.ExcelWriter(f'{file_name}') as writer:
            # Write data to Excel file
            data.to_excel(writer, index=False)

            # Show a success message
            QMessageBox.information(parent, "Sucesso", f"Os dados foram salvos com sucesso em {file_name}")
            
    except ValueError as e:
        QMessageBox.critical(parent, "Erro", str(e))

def generate_report(symbol, file_name):
    """
    Generate a stock report and save it as a DOCX file.

    Parameters
    ----------
    symbol : str
        Stock symbol
    start_date : str
        Start date in DD/MM/YYYY format
    end_date : str
        End date in DD/MM/YYYY format
    file_name : str
        File name to be saved
    """
    try:
        start_date = (datetime.datetime.now() - datetime.timedelta(days=365)).strftime('%Y-%m-%d')
        end_date = datetime.datetime.now().strftime('%Y-%m-%d')
        
        data = _fetch(symbol, start_date=start_date, end_date=end_date)

        # Rename columns
        data.columns = ['Abertura', 'Alta', 'Baixa', 'Fechamento', 'Volume', 'Dividendos', 'Splits']

        # Round prices to two decimal places
        data = data.round({'Abertura': 2, 'Alta': 2, 'Baixa': 2, 'Fechamento': 2})

        # Create a new Document
        doc = docx.Document()
        doc.add_heading(f'Relatório para {symbol}', 0)

        # Add report generation date (BR format)
        doc.add_paragraph(
        f'Data do Relatório: {datetime.datetime.now().strftime("%d/%m/%Y %H:%M:%S")}\n'
        )

        # Add financial indicators
        doc.add_heading('Principais Indicadores Financeiros', level=1)
        doc.add_paragraph(f'P/VP: {fetch_pvp(symbol)}')
        doc.add_paragraph(f'P/L: {fetch_pe(symbol)}')
        doc.add_paragraph(f'ROE: {fetch_roe(symbol)}')
        doc.add_paragraph(f'Dividend Yield: {fetch_dividend_yield(symbol)}')
        doc.add_paragraph(f'Dívida/EBITDA: {fetch_debt_to_ebitda(symbol)}')
        doc.add_paragraph(f'Margem Líquida: {fetch_net_margin(symbol)}')

        doc.add_heading('Métricas Inteligentes', level=1)
        valuation_status = calculate_valuation_status(symbol)
        doc.add_paragraph(f'Status de Valorização: {valuation_status}')
        potencial_crescimento, taxa_confianca = decidir_potencial_de_crescimento(symbol)
        doc.add_paragraph(f'Potencial de Crescimento Avaliado Automaticamente: {potencial_crescimento} ({taxa_confianca:.2f}% de confiança)')

        # Add summary statistics
        media = data['Fechamento'].mean()
        maxima = data['Fechamento'].max()
        minima = data['Fechamento'].min()
        variacao_percentual = ((data["Fechamento"].iloc[-1] - data["Fechamento"].iloc[0]) / data["Fechamento"].iloc[0]) * 100
        media_volume = data['Volume'].mean()
        maxima_volume = data['Volume'].max()
        minima_volume = data['Volume'].min()
        variacao_percentual_volume = ((data['Volume'].iloc[-1] - data['Volume'].iloc[0]) / data['Volume'].iloc[0]) * 100

        doc.add_heading('Estatísticas Gerais (de 1 ano até hoje)', level=1)
        doc.add_paragraph(f'Média: {media:.2f}')
        doc.add_paragraph(f'Máxima: {maxima:.2f}')
        doc.add_paragraph(f'Mínima: {minima:.2f}')
        doc.add_paragraph(f'Variação Percentual: {variacao_percentual:.2f}%')
        doc.add_paragraph(f'Média (Volume): {media_volume:.2f}')
        doc.add_paragraph(f'Máxima (Volume): {maxima_volume:.2f}')
        doc.add_paragraph(f'Mínima (Volume): {minima_volume:.2f}')
        doc.add_paragraph(f'Variação Percentual (Volume): {variacao_percentual_volume:.2f}%')

        # Add IFR section
        current_ifr = calculate_ifr(symbol)  

        # Determine status
        if current_ifr > 70:
            status = "Sobrecompra"
            color = "red"
        elif current_ifr < 30:
            status = "Sobrevenda"
            color = "green"
        else:
            status = "Neutro"
            color = "black"

        doc.add_page_break()
        doc.add_heading('Índice de Força Relativa (IFR período 14)', level=1)
        doc.add_paragraph(f'Valor Atual: {current_ifr:.2f}')

        # Create a status paragraph
        status_paragraph = doc.add_paragraph('Status: ')  
        # Add the status as a separate run
        status_run = status_paragraph.add_run(status)  

        # Save the document
        doc.save(file_name)

        # Show a success message
        QMessageBox.information(None, "Sucesso", "O relatório foi gerado e salvo com sucesso.")

    except Exception as e:
        raise ValueError(f"Erro gerando relatório: {str(e)}")
    
def fetch_pvp(symbol: str):
    """
    Calculate the price-to-Book ratio of a ticker
    """
    ticker = symbol + ".SA" if not symbol.endswith(".SA") else symbol
    data = yf.Ticker(ticker).info
    try:
        pvp = data['priceToBook']
    except:
        QMessageBox.warning(None, "Erro", f"Erro ao obter o P/VP para {symbol}.")
        return "Erro ao obter o P/VP"
    
    return f"{pvp:.2f}"
    

def fetch_pe(symbol: str):
    ticker = symbol + ".SA" if not symbol.endswith(".SA") else symbol
    data = yf.Ticker(ticker).info
    value = data.get('trailingPE', 'Não disponível')
    if value != 'Não disponível':
        value = f"{value:.2f}"
    return value

def fetch_roe(symbol: str):
    ticker = symbol + ".SA" if not symbol.endswith(".SA") else symbol
    data = yf.Ticker(ticker).info
    value = data.get('returnOnEquity', 'Não disponível')
    if value != 'Não disponível':
        value = f"{value * 100:.2f}%"
    return value

def fetch_dividend_yield(symbol: str):
    ticker = symbol + ".SA" if not symbol.endswith(".SA") else symbol
    data = yf.Ticker(ticker).info
    value = data.get('dividendYield', 'Não disponível')
    if isinstance(value, (int, float)):
        value = f"{value * 100:.2f}%"
    return value

import yfinance as yf

def fetch_debt_to_ebitda(symbol: str):
    # Carregar o objeto da empresa
    empresa = yf.Ticker(symbol)
    
    # Obter as demonstrações financeiras
    financials = empresa.financials
    
    # Obter a Dívida Total (Short + Long Term Debt)
    # Isso pode ser obtido do balanço patrimonial
    balance_sheet = empresa.balance_sheet
    divida_total = balance_sheet.loc['Total Debt'].iloc[0] if 'Total Debt' in balance_sheet.index else 0
    
    # Obter o Lucro Operacional (EBIT) - normalmente encontrado na Demonstração de Resultados
    lucro_operacional = financials.loc['EBIT'].iloc[0] if 'EBIT' in financials.index else 0
    
    # Obter Depreciação e Amortização
    # A depreciação e amortização pode ser obtida do fluxo de caixa (Cash Flow Statement)
    cashflow = empresa.cashflow
    depreciacao = cashflow.loc['Depreciation'].iloc[0] if 'Depreciation' in cashflow.index else 0
    amortizacao = cashflow.loc['Amortization'].iloc[0] if 'Amortization' in cashflow.index else 0
    
    # Calcular o EBITDA
    ebitda = lucro_operacional + depreciacao + amortizacao
    
    # Calcular Dívida/EBITDA
    divida_ebitda = divida_total / ebitda if ebitda else None
    
    return divida_ebitda

def fetch_net_margin(symbol: str):
    ticker = symbol + ".SA" if not symbol.endswith(".SA") else symbol
    data = yf.Ticker(ticker).info
    value = data.get('profitMargins', 'Não disponível')
    if value != 'Não disponível':
        value = f"{value * 100:.2f}%"
    return value

def get_historical_prices(symbol, period):
    # Converte o período numérico para o formato aceito pelo yfinance
    period_mapping = {
        1: '1d',
        5: '5d',
        30: '1mo',
        90: '3mo',
        180: '6mo',
        365: '1y'
    }
    period_str = period_mapping.get(period, '1mo')  # Default para '1mo' se o período não estiver mapeado
    data = fetch(symbol, period_str)
    return data

def fetch(symbol, period=None, start_date=None, end_date=None):
    symbol = symbol.upper()
    if not symbol.endswith('.SA'):
        symbol += '.SA'

    try:
        stock = yf.Ticker(symbol)
        if start_date and end_date:
            data = stock.history(start=start_date, end=end_date)
        else:
            data = stock.history(period=period)

        if data.empty:
            raise ValueError("No data available for this stock symbol")

        return data
    except Exception as e:
        raise ValueError(f"Error fetching data: {str(e)}")

def convert_to_brl_naturallanguage(value: float) -> str:
    # Convert to Brazilian natural language, like 52 bilhões instead of 52000000000

    def _format_number(number: float) -> str:
        """
        Format number to Brazilian natural language, like 52 bilhões instead of 52000000000
        """
        if not isinstance(number, (int, float)):
            return str(number)
        if number > 0:
            if number < 1000:
                return f"{number:.2f}"
            elif number < 1_000_000:
                return f"{number / 1_000:.2f} mil"
            elif number < 1_000_000_000:
                return f"{number / 1_000_000:.2f} milhões"
            elif number < 1_000_000_000_000:
                return f"{number / 1_000_000_000:.2f} bilhões"
            else:
                return f"{number / 1_000_000_000_000:.2f} trilhões"
        else:
            if number > -1000:
                return f"{number:.2f}"
            elif number > -1_000_000:
                return f"{number / -1_000:.2f} mil"
            elif number > -1_000_000_000:
                return f"{number / -1_000_000:.2f} milhões"
            elif number > -1_000_000_000_000:
                return f"{number / -1_000_000_000:.2f} bilhões"
            else:
                return f"{number / -1_000_000_000_000:.2f} trilhões"

    return _format_number(value)

def fetch_stock_data(ticker):
    ticker = yf.Ticker(ticker)
    data = ticker.info

    # Retrieve key financial data
    shares_outstanding = data.get('sharesOutstanding', 'Indeterminado')

    # Attempt to fetch balance sheet
    try:
        balance_sheet = ticker.balance_sheet
        equity = balance_sheet.loc['Total Stockholder Equity'].iloc[0] if 'Total Stockholder Equity' in balance_sheet.index else 'Indeterminado'
        total_liabilities = balance_sheet.loc['Total Liab'].iloc[0] if 'Total Liab' in balance_sheet.index else 'Indeterminado'
        total_assets = balance_sheet.loc['Total Assets'].iloc[0] if 'Total Assets' in balance_sheet.index else 'Indeterminado'
        cash = balance_sheet.loc['Cash'].iloc[0] if 'Cash' in balance_sheet.index else 'Indeterminado'
        short_term_investments = balance_sheet.loc['Short Term Investments'].iloc[0] if 'Short Term Investments' in balance_sheet.index else 'Indeterminado'
    except Exception as e:
        equity = total_liabilities = total_assets = cash = short_term_investments = 'Não disponível'
        print(f"Error fetching balance sheet: {e}")

    try:
        previous_close = data['previousClose']
    except KeyError:
        previous_close = 'Indeterminado'

    try:
        current_price = data['currentPrice']
    except KeyError:
        current_price = 'Indeterminado'

    # Calculate price ranges for today's variation
    if isinstance(previous_close, (int, float)) and isinstance(current_price, (int, float)):
        today_variation = f"R$ {min(previous_close, current_price):.2f} - R$ {max(previous_close, current_price):.2f}"
    else:
        today_variation = 'Indeterminado'

    # Calculate price ranges for yearly variation
    try:
        stock = ticker
        history = stock.history(period='1y')
        yearly_low = history['Low'].min()
        yearly_high = ticker.info['fiftyTwoWeekHigh']
        year_variation = f"R$ {min(yearly_low, yearly_high):.2f} - R$ {max(yearly_low, yearly_high):.2f}"
    except KeyError:
        year_variation = 'Indeterminado'

    try:
        last_close = data['previousClose']
    except KeyError:
        last_close = 'Indeterminado'

    try:
        market_cap = data['marketCap']
    except KeyError:
        market_cap = 'Indeterminado'

    try:
        average_volume = data['averageVolume']
    except KeyError:
        average_volume = 'Indeterminado'

    try:
        dividend_yield = data['dividendYield']
    except KeyError:
        dividend_yield = 'Indeterminado'

    try:
        main_exchange = data['exchange']
    except KeyError:
        main_exchange = 'Indeterminado'

    last_close = f"R$ {last_close:.2f}"
    market_cap = convert_to_brl_naturallanguage(market_cap)
    average_volume = convert_to_brl_naturallanguage(average_volume)
    dividend_yield = dividend_yield
    main_exchange = main_exchange

    return {
        'sharesOutstanding': shares_outstanding,
        'equity': equity,
        'totalLiabilities': total_liabilities,
        'totalAssets': total_assets,
        'cash': cash,
        'shortTermInvestments': short_term_investments,
        'last_close': last_close,
        'today_variation': today_variation,
        'year_variation': year_variation,
        'market_cap': market_cap,
        'average_volume': average_volume,
        'dividend_yield': dividend_yield,
        'main_exchange': main_exchange
    }


def fetch_monthly_financials(symbol: str):
    ticker = yf.Ticker(symbol)
    data = ticker.financials

    # Extract the relevant financial metrics
    try:
        revenue = data.loc['Total Revenue'].iloc[0]
    except KeyError:
        revenue = 'Não disponível'

    try:
        operating_expenses = data.loc['Total Operating Expenses'].iloc[0]
    except KeyError:
        operating_expenses = 'Não disponível'

    try:
        net_income = data.loc['Net Income'].iloc[0]
    except KeyError:
        net_income = 'Não disponível'

    # Calculate net profit margin
    net_profit_margin = (net_income / revenue * 100) if revenue != 0 else 'Não disponível'
    net_profit_margin = f"{net_profit_margin:.2f}%" if isinstance(net_profit_margin, float) else net_profit_margin

    try:
        earnings_per_share = ticker.info.get('trailingEps', 'Não disponível')
    except KeyError:
        earnings_per_share = 'Não disponível'

    try:
        ebitda = data.loc['EBITDA'].iloc[0]
    except KeyError:
        ebitda = 'Não disponível'

    # Effective tax rate calculation would require tax and pretax income
    try:
        tax_expenses = data.loc['Income Tax Expense'].iloc[0]
        pretax_income = net_income + tax_expenses  # Assuming net income is after tax
        effective_tax_rate = (tax_expenses / pretax_income) * 100 if pretax_income != 0 else 'Não disponível'
        effective_tax_rate = f"{effective_tax_rate:.2f}%" if isinstance(effective_tax_rate, float) else effective_tax_rate
    except KeyError:
        effective_tax_rate = 'Não disponível'

    # Additional metrics
    try:
        research_and_development_expenses = data.loc['Research Development'].iloc[0]
    except KeyError:
        research_and_development_expenses = 'Não disponível'

    try:
        selling_general_and_administrative_expenses = data.loc['Selling General Administrative'].iloc[0]
    except KeyError:
        selling_general_and_administrative_expenses = 'Não disponível'

    try:
        depreciation_expenses = ticker.cashflow.loc['Depreciation'].iloc[0]
    except KeyError:
        depreciation_expenses = 'Não disponível'

    try:
        interest_expenses = data.loc['Interest Expense'].iloc[0]
    except KeyError:
        interest_expenses = 'Não disponível'

    try:
        total_expenses = data.loc['Total Expenses'].iloc[0]
    except KeyError:
        total_expenses = 'Não disponível'

    # Return all metrics
    return {
        'revenue': convert_to_brl_naturallanguage(revenue),
        'operating_expenses': operating_expenses,
        'net_income': net_income,
        'net_profit_margin': net_profit_margin,
        'earnings_per_share': earnings_per_share,
        'ebitda': ebitda,
        'effective_tax_rate': effective_tax_rate,
        'research_and_development_expenses': research_and_development_expenses,
        'selling_general_and_administrative_expenses': selling_general_and_administrative_expenses,
        'depreciation_expenses': depreciation_expenses,
        'interest_expenses': interest_expenses,
        'total_expenses': total_expenses,
    }

def fetch_quarterly_data(ticker, num_quarters=4):
    """
    Fetch the specified number of quarterly total assets and total liabilities for the given ticker.
    """
    try:
        data = yf.Ticker(ticker).financials
        #print("Data fetched for quarterly:", data)  # Debug print
        total_assets = data.loc['Total Assets'].iloc[:, :num_quarters]  # Last num_quarters
        total_liabilities = data.loc['Total Liabilities Net'].iloc[:, :num_quarters]  # Last num_quarters
        return total_assets, total_liabilities
    except KeyError as e:
        #print(f"Error fetching quarterly data: {e}")
        return [0] * num_quarters, [0] * num_quarters
    except Exception as e:
        #print(f"Unexpected error fetching quarterly data: {e}")
        return [0] * num_quarters, [0] * num_quarters


def fetch_annual_data(ticker, num_years=3):
    """
    Fetch the specified number of annual total assets and total liabilities for the given ticker.
    """
    try:
        data = yf.Ticker(ticker).financials
        #print("Data fetched for annual:", data)  # Debug print
        total_assets = data.loc['Total Assets'].iloc[:, :num_years]  # Last num_years
        total_liabilities = data.loc['Total Liabilities Net'].iloc[:, :num_years]  # Last num_years
        return total_assets, total_liabilities
    except KeyError as e:
        #print(f"Error fetching annual data: {e}")
        return [0] * num_years, [0] * num_years
    except Exception as e:
        #print(f"Unexpected error fetching annual data: {e}")
        return [0] * num_years, [0] * num_years

def calculate_ifr(symbol, period=14):
    # Fetch historical data for the last 365 days
    historical_data = _fetch(symbol, start_date=(datetime.datetime.now() - datetime.timedelta(days=365)).strftime('%Y-%m-%d'), end_date=datetime.datetime.now().strftime('%Y-%m-%d'))

    # Calculate daily price changes
    delta = historical_data['Close'].diff()
    
    # Separate gains and losses
    gains = (delta.where(delta > 0, 0)).rolling(window=period).mean()
    losses = (-delta.where(delta < 0, 0)).rolling(window=period).mean()
    
    # Calculate Relative Strength (RS)
    rs = gains / losses
    
    # Calculate IFR
    if losses.iloc[-1] == 0:  # Avoid division by zero
        return 100  # If there are no losses, the IFR is 100
    else:
        ifr = 100 - (100 / (1 + rs.iloc[-1]))
        return ifr

def calculate_macd(prices, short_window=12, long_window=26, signal_window=9):
    # Calculate short-term EMA
    short_ema = prices.ewm(span=short_window, adjust=False).mean()
    # Calculate long-term EMA
    long_ema = prices.ewm(span=long_window, adjust=False).mean()
    # Calculate MACD line
    macd = short_ema - long_ema
    # Calculate signal line
    signal = macd.ewm(span=signal_window, adjust=False).mean()
    return macd, signal

def calcular_media_movel(precos, periodo):
    return pd.Series(precos).rolling(window=periodo).mean()

def calcular_desvio_padrao(precos, periodo):
    return pd.Series(precos).rolling(window=periodo).std()

def calcular_bandas_bollinger(precos, periodo):
    media_movel = calcular_media_movel(precos, periodo)
    desvio_padrao = calcular_desvio_padrao(precos, periodo)
    banda_superior = media_movel + (desvio_padrao * 2)
    banda_inferior = media_movel - (desvio_padrao * 2)
    return media_movel, banda_superior, banda_inferior

def calcular_estocastico_normal(precos, n=14):
    L_n = precos['Low'].rolling(window=n).min()  # Menor preço em n período
    H_n = precos['High'].rolling(window=n).max()  # Maior preço em n período
    C = precos['Close']  # Preço de fechamento mais recente

    K = 100 * ((C - L_n) / (H_n - L_n))  # Cálculo do Estocástico K
    return K

def calcular_estocastico_lento(precos, n=14, d=3):
    K = calcular_estocastico_normal(precos, n)
    D = K.rolling(window=d).mean()  # Média móvel do K
    return K, D

def calculate_valuation_status(ticker: str):
    try:
        pvp = float(fetch_pvp(ticker))  # Use the ticker to get the P/VP
    except TypeError:
        return "Indisponível"

    if pvp < 0.95:
        return "Subvalorizado"
    elif 0.95 <= pvp <= 1.05:
        return "Balanceado"
    else:
        return "Supervalorizado"

def avaliar_indicador(indicador, tipo='positivo'):
    if indicador is None:
        return 'Indeterminado'
    
    # Para indicadores onde maior é melhor
    if tipo == 'positivo': 
        if indicador >= 0.10:  # Crescimento e outros indicadores positivos
            return 'Acima da média'
        elif indicador >= 0.05:
            return 'Regular'
        else:
            return 'Abaixo da média'
    
    # Para PEG Ratio
    elif tipo == 'peg_ratio':
        if indicador < 1.0:  # PEG abaixo de 1, é considerado bom
            return 'Acima da média'
        elif indicador < 2.0:
            return 'Regular'
        else:
            return 'Abaixo da média'

    # Para margens, ROIC e outros onde maior é melhor
    elif tipo == 'margem':  
        if indicador >= 0.30:  # Margens mais altas agora consideradas "Acima da média"
            return 'Acima da média'
        elif indicador >= 0.20:
            return 'Regular'
        else:
            return 'Abaixo da média'

    # Para ROIC
    elif tipo == 'roic':  
        if indicador >= 0.15:  # ROIC acima de 15%
            return 'Acima da média'
        elif indicador >= 0.08:
            return 'Regular'
        else:
            return 'Abaixo da média'

    # Para indicadores onde menor é melhor (como Dívida)
    elif tipo == 'negativo':  
        if indicador < 0.5:  # Menor dívida é melhor
            return 'Acima da média'
        elif indicador < 1:
            return 'Regular'
        else:
            return 'Abaixo da média'

def decidir_potencial_de_crescimento(ticker):
    # Baixar dados financeiros do ticker
    ticker = ticker.upper()
    if not ticker.endswith('.SA'):
        ticker += '.SA'
    
    empresa = yf.Ticker(ticker)
    
    # Pegar os dados financeiros principais
    info = empresa.info
    
    # Indicadores de Crescimento
    eps_growth = info.get('earningsQuarterlyGrowth', None)  # Crescimento do EPS
    peg_ratio = info.get('trailingPE', None) / eps_growth if eps_growth else None  # PEG Ratio
    receita_crescimento = None
    if info.get('regularMarketPreviousClose') and info.get('marketCap'):
        receita_crescimento = (info.get('regularMarketPreviousClose', None) - info.get('marketCap', 0)) / info.get('marketCap', 1)  # Crescimento da Receita
    margem_bruta = info.get('grossMargins', None)  # Margem Bruta
    margem_operacional = info.get('operatingMargins', None)  # Margem Operacional
    roic = info.get('returnOnEquity', None)  # ROIC
    debt_to_equity = info.get('debtToEquity', None)  # Dívida sobre Patrimônio

    # Avaliando cada indicador separadamente usando as funções individualizadas
    resultados = {
        'EPS Growth': avaliar_eps_growth(eps_growth),
        'PEG Ratio': avaliar_peg_ratio(peg_ratio),
        'Crescimento de Receita (CAGR)': avaliar_crescimento_receita(receita_crescimento),
        'Margem Bruta': avaliar_margem_bruta(margem_bruta),
        'Margem Operacional': avaliar_margem_operacional(margem_operacional),
        'ROIC': avaliar_roic(roic),
        'Debt-to-Equity': avaliar_debt_to_equity(debt_to_equity),
    }
    
    # Imprimindo os indicadores brutos
    #print({
    #    'EPS Growth': eps_growth,
    #    'PEG Ratio': peg_ratio,
    #    'Crescimento de Receita (CAGR)': receita_crescimento,
    #    'Margem Bruta': margem_bruta,
    #    'Margem Operacional': margem_operacional,
    #    'ROIC': roic,
    #    'Debt-to-Equity': debt_to_equity,
    #})
    
    # Imprimindo os resultados avaliados
    # print(resultados)

    # Excluindo indicadores indeterminados
    votos_validos = [voto for voto in resultados.values() if voto != 'Indeterminado']
    
    if not votos_validos:  # Se não houver indicadores válidos
        return "Nenhum dado válido para análise", 0.0
    
    # Contando os votos válidos
    contagem_votos = Counter(votos_validos)
    
    # Identificando o resultado com maior voto
    resultado_final = contagem_votos.most_common(1)[0][0]
    taxa_de_confianca = (contagem_votos[resultado_final] / len(votos_validos)) * 100
    
    return resultado_final, round(taxa_de_confianca, 2)

def avaliar_eps_growth(indicador):
    if indicador is None:
        return 'Indeterminado'
    if indicador >= 0.10:
        return 'Acima da média'
    elif indicador >= 0.05:
        return 'Regular'
    else:
        return 'Abaixo da média'

def avaliar_peg_ratio(indicador):
    if indicador is None:
        return 'Indeterminado'
    if indicador < 1.0:
        return 'Acima da média'
    elif indicador < 2.0:
        return 'Regular'
    else:
        return 'Abaixo da média'

def avaliar_crescimento_receita(indicador):
    if indicador is None:
        return 'Indeterminado'
    if indicador >= 0.10:
        return 'Acima da média'
    elif indicador >= 0.05:
        return 'Regular'
    else:
        return 'Abaixo da média'

def avaliar_margem_bruta(indicador):
    if indicador is None:
        return 'Indeterminado'
    if indicador >= 0.30:
        return 'Acima da média'
    elif indicador >= 0.20:
        return 'Regular'
    else:
        return 'Abaixo da média'

def avaliar_margem_operacional(indicador):
    if indicador is None:
        return 'Indeterminado'
    if indicador >= 0.30:
        return 'Acima da média'
    elif indicador >= 0.20:
        return 'Regular'
    else:
        return 'Abaixo da média'

def avaliar_roic(indicador):
    if indicador is None:
        return 'Indeterminado'
    if indicador >= 0.15:
        return 'Acima da média'
    elif indicador >= 0.08:
        return 'Regular'
    else:
        return 'Abaixo da média'

def avaliar_debt_to_equity(indicador):
    if indicador is None:
        return 'Indeterminado'
    if indicador < 0.5:
        return 'Acima da média'
    elif indicador < 1:
        return 'Regular'
    else:
        return 'Abaixo da média'