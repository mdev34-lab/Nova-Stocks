import sys
from PySide6.QtWidgets import (
    QApplication, QMainWindow, QMenu, QMenuBar, QVBoxLayout, QWidget, QPushButton, QFrame, QInputDialog, QSizePolicy, QMessageBox, QDateEdit, QDialog, QDialogButtonBox, QLabel, QFileDialog, QTableWidget, QTableWidgetItem, QDockWidget
)
from PySide6.QtGui import QAction
from PySide6.QtCore import Qt, QThread, Signal, QDate, QTimer
import matplotlib.pyplot as plt
from matplotlib.backends.backend_qt5agg import FigureCanvasQTAgg as FigureCanvas
import pandas as pd
from datetime import datetime, date, timedelta
import numpy as np
from stocklibs import stockdata
from stocklibs.stockdata import calculate_macd, calcular_media_movel, calcular_desvio_padrao, calcular_bandas_bollinger, calcular_estocastico_normal, calcular_estocastico_lento
from stocklibs.metrics import MetricsWindow
from stocklibs.smart_metrics import SmartMetricsWindow
from stocklibs.revenue_income_chart import RevenueIncomeChart
from stocklibs.assets_liabilities_chart import AssetsLiabilitiesChart
from stocklibs.analysis import StockAnalysis
from stocklibs.assets import styles
from stocklibs.settings_dialog import SettingsDialog, SettingsManager
import matplotlib.gridspec as gridspec
import matplotlib.dates as mdates
import yfinance as yf
import os

# Constantes
DEFAULT_CANDLESTICK_PERIOD = 1
DEFAULT_START_DATE = date.today() - timedelta(days=365)
DEFAULT_END_DATE = date.today()

class DataFetcher:
    @staticmethod
    def fetch_stock_data(symbol, start_date, end_date, candlestick_period, cache):
        # Convert QDate to string format if necessary
        if isinstance(start_date, QDate):
            start_date = start_date.toString("yyyy-MM-dd")
        if isinstance(end_date, QDate):
            end_date = end_date.toString("yyyy-MM-dd")

        cache_key = (symbol, start_date, end_date, candlestick_period)
        if cache_key in cache:
            return cache[cache_key]
        try:
            data = stockdata.fetch(symbol=symbol, start_date=start_date, end_date=end_date)
            if data.empty or data.isnull().values.any():
                raise ValueError(f"Não há dados válidos para {symbol}")
            data = data[['Open', 'High', 'Low', 'Close', 'Volume']].dropna()
            if 'Volume' not in data.columns:
                raise ValueError("Dados de volume não disponíveis")
            if candlestick_period > 1:
                data['group'] = np.arange(len(data)) // candlestick_period
                data = pd.DataFrame({
                    'Open': data.groupby('group')['Open'].first(),
                    'High': data.groupby('group')['High'].max(),
                    'Low': data.groupby('group')['Low'].min(),
                    'Close': data.groupby('group')['Close'].last(),
                    'Volume': data.groupby('group')['Volume'].sum(),
                    'Date': data.groupby('group').apply(lambda x: x.index[0])
                })
                data.set_index('Date', inplace=True)
            cache[cache_key] = data
            return data
        except ValueError as e:
            raise ValueError(f"Erro ao buscar dados para {symbol}: {e}")

class Plotter:
    @staticmethod
    def plot_candlestick_chart(canvas, data, ticker, settings, medias, show_volume, show_ifr, show_macd, show_bandas_bollinger, show_estocastico_normal, show_estocastico_lento, candlestick_period):
        canvas.figure.clear()
        canvas.figure.patch.set_facecolor('#252525')  # Set background color
        if show_ifr:
            gs = gridspec.GridSpec(2, 1, height_ratios=[3, 1])
            ax1 = canvas.figure.add_subplot(gs[0])
            ax2 = canvas.figure.add_subplot(gs[1])
        elif show_volume:
            gs = gridspec.GridSpec(2, 1, height_ratios=[3, 1])
            ax1 = canvas.figure.add_subplot(gs[0])
            ax3 = canvas.figure.add_subplot(gs[1])
        elif show_macd:
            gs = gridspec.GridSpec(2, 1, height_ratios=[3, 1])
            ax1 = canvas.figure.add_subplot(gs[0])
            ax4 = canvas.figure.add_subplot(gs[1])
        elif show_estocastico_normal:
            gs = gridspec.GridSpec(2, 1, height_ratios=[3, 1])
            ax1 = canvas.figure.add_subplot(gs[0])
            ax5 = canvas.figure.add_subplot(gs[1])
        elif show_estocastico_lento:
            gs = gridspec.GridSpec(2, 1, height_ratios=[3, 1])
            ax1 = canvas.figure.add_subplot(gs[0])
            ax6 = canvas.figure.add_subplot(gs[1])
        else:
            gs = gridspec.GridSpec(1, 1)
            ax1 = canvas.figure.add_subplot(gs[0])

        ax1.set_facecolor('#1e1e1e')  # Set subplot background color
        ax1.tick_params(axis='x', colors='#d4d4d4')  # Set x-axis tick color
        ax1.tick_params(axis='y', colors='#d4d4d4')  # Set y-axis tick color
        ax1.set_title(f"Gráfico de Candlestick para {ticker} (Período: {candlestick_period} {'dia' if candlestick_period == 1 else 'dias'})", color='#d4d4d4')
        ax1.set_ylabel("Preço", color='#d4d4d4')
        ax1.spines['bottom'].set_color('#333')
        ax1.spines['top'].set_color('#333')
        ax1.spines['right'].set_color('#333')
        ax1.spines['left'].set_color('#333')

        width = 0.6 * candlestick_period
        width2 = 0.05 * candlestick_period

        up = data[data.Close >= data.Open]
        down = data[data.Close < data.Open]

        if show_ifr:
            delta = data['Close'].diff()
            gain = (delta.where(delta > 0, 0)).rolling(window=settings.rsi_period).mean()
            loss = (-delta.where(delta < 0, 0)).rolling(window=settings.rsi_period).mean()
            rs = gain / loss
            rsi = 100 - (100 / (1 + rs))
            ax2.set_facecolor('#1e1e1e')
            ax2.tick_params(axis='x', colors='#d4d4d4')
            ax2.tick_params(axis='y', colors='#d4d4d4')
            ax2.set_title("IFR", color='#d4d4d4')
            ax2.set_ylabel("Valor do IFR", color='#d4d4d4')
            ax2.plot(data.index, rsi, color='purple', label='IFR')
            ax2.axhline(30, color='red', linestyle='--', label='SV')
            ax2.axhline(70, color='green', linestyle='--', label='SC')
            ax2.set_ylim(0, 100)
            ax2.legend(labelcolor='#d4d4d4')
            ax2.spines['bottom'].set_color('#333')
            ax2.spines['top'].set_color('#333')
            ax2.spines['right'].set_color('#333')
            ax2.spines['left'].set_color('#333')
        elif show_volume:
            ax3.set_facecolor('#1e1e1e')
            ax3.tick_params(axis='x', colors='#d4d4d4')
            ax3.tick_params(axis='y', colors='#d4d4d4')
            ax3.set_title("Volume", color='#d4d4d4')
            ax3.set_ylabel("Volume", color='#d4d4d4')
            ax3.bar(data.index, data['Volume'], color='b', alpha=0.5)
            ax3.set_ylim(0, data['Volume'].max() * 1.1)
            ax3.spines['bottom'].set_color('#333')
            ax3.spines['top'].set_color('#333')
            ax3.spines['right'].set_color('#333')
            ax3.spines['left'].set_color('#333')
        elif show_macd:
            macd, signal = calculate_macd(data['Close'], settings.macd_fast_period, settings.macd_slow_period, settings.macd_signal_period)
            ax4.set_facecolor('#1e1e1e')
            ax4.tick_params(axis='x', colors='#d4d4d4')
            ax4.tick_params(axis='y', colors='#d4d4d4')
            ax4.set_title("MACD", color='#d4d4d4')
            ax4.set_ylabel("MACD", color='#d4d4d4')
            ax4.plot(data.index, macd, color='blue', label='MACD')
            ax4.plot(data.index, signal, color='red', label='Sinal')
            histogram = macd - signal
            ax4.bar(data.index, histogram, color='gray', label='Histograma')
            ax4.axhline(0, color='black', linestyle='--')
            ax4.legend(labelcolor='#d4d4d4')
            ax4.spines['bottom'].set_color('#333')
            ax4.spines['top'].set_color('#333')
            ax4.spines['right'].set_color('#333')
            ax4.spines['left'].set_color('#333')
        elif show_estocastico_normal:
            estocastico = calcular_estocastico_normal(data, settings.stochastic_k_period)
            ax5.set_facecolor('#1e1e1e')
            ax5.tick_params(axis='x', colors='#d4d4d4')
            ax5.tick_params(axis='y', colors='#d4d4d4')
            ax5.set_title("Estocástico Normal", color='#d4d4d4')
            ax5.set_ylabel("Valor Estocástico", color='#d4d4d4')
            ax5.plot(estocastico, label='Estocástico Normal', color='purple')
            ax5.axhline(20, color='red', linestyle='--', label='SV')
            ax5.axhline(80, color='green', linestyle='--', label='SC')
            ax5.legend(labelcolor='#d4d4d4')
            ax5.spines['bottom'].set_color('#333')
            ax5.spines['top'].set_color('#333')
            ax5.spines['right'].set_color('#333')
            ax5.spines['left'].set_color('#333')
        elif show_estocastico_lento:
            estocastico_lento, estocastico_d = calcular_estocastico_lento(data, settings.stochastic_k_period, settings.stochastic_d_period)
            ax6.set_facecolor('#1e1e1e')
            ax6.tick_params(axis='x', colors='#d4d4d4')
            ax6.tick_params(axis='y', colors='#d4d4d4')
            ax6.set_title("Estocástico Lento", color='#d4d4d4')
            ax6.set_ylabel("Valor Estocástico", color='#d4d4d4')
            ax6.plot(estocastico_lento, label='K Estocástico Lento', color='purple')
            ax6.plot(estocastico_d, label='D Estocástico Lento', color='orange')
            ax6.axhline(20, color='red', linestyle='--', label='SV')
            ax6.axhline(80, color='green', linestyle='--', label='SC')
            ax6.legend(labelcolor='#d4d4d4')
            ax6.spines['bottom'].set_color('#333')
            ax6.spines['top'].set_color('#333')
            ax6.spines['right'].set_color('#333')
            ax6.spines['left'].set_color('#333')

        if show_bandas_bollinger:
            media_movel, banda_superior, banda_inferior = calcular_bandas_bollinger(data['Close'], settings.ma_period)
            ax1.plot(media_movel, label='Média Móvel', color='blue')
            ax1.plot(banda_superior, label='Banda Superior', color='red', linestyle='--')
            ax1.plot(banda_inferior, label='Banda Inferior', color='green', linestyle='--')
            ax1.fill_between(data.index, banda_inferior, banda_superior, color='gray', alpha=0.3)
            ax1.legend(labelcolor='#d4d4d4', facecolor='#1e1e1e', edgecolor='#333')

        plt.tight_layout()

        ax1.bar(up.index, up.Close - up.Open, width, bottom=up.Open, color='g')
        ax1.bar(up.index, up.High - up.Close, width2, bottom=up.Close, color='g')
        ax1.bar(up.index, up.Low - up.Open, width2, bottom=up.Open, color='g')

        ax1.bar(down.index, down.Close - down.Open, width, bottom=down.Open, color='r')
        ax1.bar(down.index, down.High - down.Open, width2, bottom=down.Open, color='r')
        ax1.bar(down.index, down.Low - down.Close, width2, bottom=down.Close, color='r')

        if 'SMA' in medias:
            sma = data['Close'].rolling(window=settings.ma_period).mean()
            ax1.plot(data.index, sma, color='cyan', label='SMA')
        if 'EMA' in medias:
            ema = data['Close'].ewm(span=settings.ema_period, adjust=False).mean()
            ax1.plot(data.index, ema, color='blue', label='EMA')
        if 'WMA' in medias:
            wma = data['Close'].rolling(window=settings.wma_period).apply(lambda x: np.average(x, weights=np.arange(len(x), 0, -1)))
            ax1.plot(data.index, wma, color='#6495ED', label='WMA')

        ax1.grid(True, color='#333')
        ax1.xaxis.set_major_locator(plt.MaxNLocator(10))
        plt.xticks(rotation=45, color='#d4d4d4')
        plt.tight_layout()
        canvas.figure.tight_layout()

        if 'SMA' in medias or 'EMA' in medias or 'WMA' in medias:
            ax1.legend(labelcolor='#d4d4d4', facecolor='#1e1e1e', edgecolor='#333')

        def zoom(event):
            if event.inaxes is not None:
                xlim = ax1.get_xlim()
                xlim_dates = [mdates.num2date(x) for x in xlim]
                mouse_x = mdates.num2date(event.xdata) if event.xdata else None
                current_range_seconds = (xlim_dates[1] - xlim_dates[0]).total_seconds()
                zoom_amount = pd.Timedelta(seconds=current_range_seconds * 0.3)
                if event.button == 'down':
                    new_xlim = [mouse_x - zoom_amount, mouse_x + zoom_amount] if mouse_x else xlim_dates
                elif event.button == 'up':
                    new_xlim = [xlim_dates[0] - zoom_amount, xlim_dates[1] + zoom_amount]
                data_start = mdates.date2num(data.index[0])
                data_end = mdates.date2num(data.index[-1])
                new_xlim_numeric = [mdates.date2num(x) for x in new_xlim]
                new_xlim_clamped = [
                    max(data_start, new_xlim_numeric[0]),
                    min(data_end, new_xlim_numeric[1])
                ]
                ax1.set_xlim(mdates.num2date(new_xlim_clamped[0]), mdates.num2date(new_xlim_clamped[1]))
                canvas.draw()

        canvas.mpl_connect('scroll_event', zoom)
        canvas.draw()

class IndicatorUpdater:
    def __init__(self, pvp_indicator, pe_indicator, roe_indicator, dividend_yield_indicator, debt_to_ebitda_indicator, net_margin_indicator):
        self.pvp_indicator = pvp_indicator
        self.pe_indicator = pe_indicator
        self.roe_indicator = roe_indicator
        self.dividend_yield_indicator = dividend_yield_indicator
        self.debt_to_ebitda_indicator = debt_to_ebitda_indicator
        self.net_margin_indicator = net_margin_indicator

    def update_indicators(self, ticker):
        try:
            ticker_pvp = stockdata.fetch_pvp(ticker.upper())
            ticker_pe = stockdata.fetch_pe(ticker.upper())
            ticker_roe = stockdata.fetch_roe(ticker.upper())
            ticker_dividend_yield = stockdata.fetch_dividend_yield(ticker.upper())
            ticker_debt_to_ebitda = stockdata.fetch_debt_to_ebitda(ticker.upper())
            ticker_net_margin = stockdata.fetch_net_margin(ticker.upper())
        except Exception:
            ticker_pvp = ticker_pe = ticker_roe = ticker_dividend_yield = ticker_debt_to_ebitda = ticker_net_margin = "Não disponível"

        self.pvp_indicator.setText(f"P/VP: <b>{ticker_pvp}</b>")
        self.pe_indicator.setText(f"P/L: <b>{ticker_pe}</b>")
        self.roe_indicator.setText(f"ROE: <b>{ticker_roe}</b>")
        self.dividend_yield_indicator.setText(f"Dividend Yield: <b>{ticker_dividend_yield}</b>")
        if ticker_debt_to_ebitda is not None:
            self.debt_to_ebitda_indicator.setText(f"Dívida/EBITDA: <b>{ticker_debt_to_ebitda:.2f}</b>")
        else:
            self.debt_to_ebitda_indicator.setText("Dívida/EBITDA: <b>Não disponível</b>")
        self.net_margin_indicator.setText(f"Margem Líquida: <b>{ticker_net_margin}</b>")

class NovaGUI(QMainWindow):
    def __init__(self):
        super().__init__()
        self.current_settings = SettingsManager()
        self.current_analysis = StockAnalysis()
        self.current_analysis.start_date = self.current_settings.start_date
        self.current_analysis.end_date = self.current_settings.end_date
        self.current_analysis.candlestick_period = self.current_settings.candlestick_period
        self.candlestick_cache = {}
        self.mouse_move_timer = QTimer(self)
        self.mouse_move_timer.setSingleShot(True)
        self.mouse_move_timer.timeout.connect(self.delayed_draw)
        self.last_mouse_event = None

        self.setWindowTitle("Nova Stocks")
        self.setGeometry(100, 100, 900, 600)
        self.setStyleSheet(styles.light_mode)

        self.central_widget = QWidget(self)
        self.setCentralWidget(self.central_widget)
        self.main_layout = QVBoxLayout(self.central_widget)

        self.graph_frame = QFrame(self)
        self.graph_frame.setStyleSheet("background-color: #252525;")
        self.graph_layout = QVBoxLayout(self.graph_frame)
        self.main_layout.addWidget(self.graph_frame)

        self.canvas = FigureCanvas(plt.Figure())
        self.canvas.figure.patch.set_facecolor('#252525')  # Set background color
        self.canvas.setSizePolicy(QSizePolicy.Expanding, QSizePolicy.Expanding)
        self.graph_layout.addWidget(self.canvas)

        self.menubar = QMenuBar(self)
        self.setMenuBar(self.menubar)

        self.file_menu = QMenu("Arquivo", self)
        self.menubar.addMenu(self.file_menu)
        self.file_menu.addAction("Nova Análise", self.nova_analise).setShortcut('Ctrl+N')
        self.file_menu.addSeparator()
        self.file_menu.addAction("Configurações", self.open_settings)
        self.file_menu.addAction("Sair", self.close)

        self.edit_menu = QMenu("Editar", self)
        self.menubar.addMenu(self.edit_menu)
        self.edit_menu.addAction("Alterar Período de Análise", self.toggle_custom_date)
        self.edit_menu.addAction("Alterar Período das Médias Móveis", self.toggle_ma_period)
        self.edit_menu.addAction("Alterar Período dos Candlesticks", self.toggle_candlestick_period)
        self.edit_menu.addSeparator()
        self.edit_menu.addAction("Desfazer", self.desfazer).setShortcut('Ctrl+Z')

        self.view_menu = QMenu("Visualizar", self)
        self.menubar.addMenu(self.view_menu)
        self.view_menu.addAction("Média Móvel Simples", self.mostrar_media_movel_simples)
        self.view_menu.addAction("Média Móvel Exponencial", self.mostrar_media_movel_exponencial)
        self.view_menu.addAction("Média Móvel Ponderada", self.mostrar_media_ponderada)
        self.view_menu.addSeparator()
        self.view_menu.addAction("Volumes", self.mostrar_volume)
        self.view_menu.addAction("Índice de Força Relativa", self.toggle_ifr)
        self.view_menu.addAction("Média Móvel de Convergência/Divergência (MACD)", self.toggle_macd)
        self.view_menu.addAction("Bandas de Bollinger", self.mostrar_bandas_bollinger)
        self.view_menu.addAction("Indicador Estocástico", self.toggle_estocastico_normal)
        self.view_menu.addAction("Indicador Estocástico Lento", self.toggle_estocastico_lento)
        self.view_menu.addSeparator()

        self.informacoes_menu = QMenu("Informações", self)
        self.informacoes_menu.addAction("Principais Métricas", self.toggle_sidebar)
        self.informacoes_menu.addAction("Todas as Métricas", self.show_detailed_metrics)
        self.informacoes_menu.addAction("Gráfico Receita/Renda Líquida", self.show_revenue_income_chart)
        self.assets_liabilities_action = QAction("Gráfico Ativos/Passivos", self)
        self.assets_liabilities_action.triggered.connect(self.show_assets_liabilities_chart)
        self.informacoes_menu.addAction(self.assets_liabilities_action)
        self.menubar.addMenu(self.informacoes_menu)

        self.estrategias_menu = QMenu("Estratégias", self)
        self.menubar.addMenu(self.estrategias_menu)
        self.smart_metrics_action = QAction("Métricas Inteligentes", self)
        self.smart_metrics_action.triggered.connect(self.open_smart_metrics)
        self.estrategias_menu.addAction(self.smart_metrics_action)

        self.export_menu = QMenu("Exportar", self)
        self.menubar.addMenu(self.export_menu)
        self.export_menu.addAction("Exportar Dados para XLSX", self.export_to_excel)
        self.intelligent_report_menu = self.export_menu.addMenu("Relatório Inteligente")
        self.intelligent_report_menu.addAction("Gerar em DOCX", self.generate_intelligent_report_docx)

        self.pvp_indicator = QLabel(f"P/VP: Indefinido")
        self.pvp_indicator.setSizePolicy(QSizePolicy(QSizePolicy.Fixed, QSizePolicy.Fixed))
        self.pe_indicator = QLabel(f"P/L: Indefinido")
        self.pe_indicator.setSizePolicy(QSizePolicy(QSizePolicy.Fixed, QSizePolicy.Fixed))
        self.roe_indicator = QLabel(f"ROE: Indefinido")
        self.roe_indicator.setSizePolicy(QSizePolicy(QSizePolicy.Fixed, QSizePolicy.Fixed))
        self.dividend_yield_indicator = QLabel(f"Dividend Yield: Indefinido")
        self.dividend_yield_indicator.setSizePolicy(QSizePolicy(QSizePolicy.Fixed, QSizePolicy.Fixed))
        self.debt_to_ebitda_indicator = QLabel(f"Dívida/EBITDA: Indefinido")
        self.debt_to_ebitda_indicator.setSizePolicy(QSizePolicy(QSizePolicy.Fixed, QSizePolicy.Fixed))
        self.net_margin_indicator = QLabel(f"Margem Líquida: Indefinido")
        self.net_margin_indicator.setSizePolicy(QSizePolicy(QSizePolicy.Fixed, QSizePolicy.Fixed))

        layout = QVBoxLayout()
        layout.addWidget(self.pvp_indicator)
        self.pvp_indicator.setTextInteractionFlags(Qt.TextSelectableByMouse)
        layout.addWidget(self.pe_indicator)
        self.pe_indicator.setTextInteractionFlags(Qt.TextSelectableByMouse)
        layout.addWidget(self.roe_indicator)
        self.roe_indicator.setTextInteractionFlags(Qt.TextSelectableByMouse)
        layout.addWidget(self.dividend_yield_indicator)
        self.dividend_yield_indicator.setTextInteractionFlags(Qt.TextSelectableByMouse)
        layout.addWidget(self.debt_to_ebitda_indicator)
        self.debt_to_ebitda_indicator.setTextInteractionFlags(Qt.TextSelectableByMouse)
        layout.addWidget(self.net_margin_indicator)
        self.net_margin_indicator.setTextInteractionFlags(Qt.TextSelectableByMouse)

        self.indicator_updater = IndicatorUpdater(
            self.pvp_indicator,
            self.pe_indicator,
            self.roe_indicator,
            self.dividend_yield_indicator,
            self.debt_to_ebitda_indicator,
            self.net_margin_indicator
        )

        self.update_menu_state()

    def open_settings(self):
        if not (self.current_analysis.ticker is None):
            settings_dialog = SettingsDialog(self, self.current_settings)
            settings_dialog.settings_changed.connect(self.plot_chart)
            settings_dialog.exec()
        else:
            QMessageBox.warning(self, "Erro", "Você precisa definir um ticker antes de abrir as configurações.")

    def toggle_ifr(self):
        self.current_analysis.toggle_ifr()
        self.plot_chart()

    def toggle_macd(self):
        self.current_analysis.toggle_macd()
        self.plot_chart()

    def toggle_estocastico_normal(self):
        self.current_analysis.toggle_estocastico_normal()
        self.plot_chart()

    def toggle_estocastico_lento(self):
        self.current_analysis.toggle_estocastico_lento()
        self.plot_chart()

    def mostrar_bandas_bollinger(self):
        self.current_analysis.mostrar_bandas_bollinger()
        self.plot_chart()

    def mostrar_volume(self):
        self.current_analysis.show_volumes()
        self.plot_chart()

    def abrir_popup_ticker(self):
        ticker, ok = QInputDialog.getText(self, "Input", "Por favor, insira o ticker da ação:")
        if ok and ticker:
            ticker = ticker.upper()
            if not stockdata.is_valid_ticker(ticker):
                QMessageBox.warning(self, "Erro", "Ticker inválido. Por favor, insira um ticker válido.")
                return
            self.set_ticker(ticker)
            self.plot_chart()

    def plot_chart(self):
        try:
            data = DataFetcher.fetch_stock_data(
                self.current_analysis.ticker,
                self.current_analysis.start_date,
                self.current_analysis.end_date,
                self.current_analysis.candlestick_period,
                self.candlestick_cache
            )
            Plotter.plot_candlestick_chart(
                self.canvas,
                data,
                self.current_analysis.ticker,
                self.current_settings,
                self.current_analysis.medias,
                self.current_analysis.show_volume,
                self.current_analysis.show_ifr,
                self.current_analysis.show_macd,
                self.current_analysis.show_bandas_bollinger,
                self.current_analysis.show_estocastico_normal,
                self.current_analysis.show_estocastico_lento,
                self.current_analysis.candlestick_period
            )
        except ValueError as e:
            QMessageBox.warning(self, "Erro", str(e))

    def mostrar_media_movel_simples(self):
        if 'SMA' in self.current_analysis.medias:
            self.current_analysis.medias.remove('SMA')
        else:
            self.current_analysis.medias.append('SMA')
        self.plot_chart()

    def mostrar_media_movel_exponencial(self):
        if 'EMA' in self.current_analysis.medias:
            self.current_analysis.medias.remove('EMA')
        else:
            self.current_analysis.medias.append('EMA')
        self.plot_chart()

    def mostrar_media_ponderada(self):
        if 'WMA' in self.current_analysis.medias:
            self.current_analysis.medias.remove('WMA')
        else:
            self.current_analysis.medias.append('WMA')
        self.plot_chart()

    def nova_analise(self):
        self.abrir_popup_ticker()

    def toggle_custom_date(self):
        dialog = QDialog(self)
        dialog.setWindowTitle("Selecionar Período de Análise")

        date_edit_start = QDateEdit(self)
        date_edit_start.setDate(self.current_analysis.start_date)
        date_edit_start.setCalendarPopup(True)

        date_edit_end = QDateEdit(self)
        date_edit_end.setDate(self.current_analysis.end_date)
        date_edit_end.setCalendarPopup(True)

        dialog_layout = QVBoxLayout()
        dialog_layout.addWidget(QLabel("Data de Início"))
        dialog_layout.addWidget(date_edit_start)
        dialog_layout.addWidget(QLabel("Data de Fim"))
        dialog_layout.addWidget(date_edit_end)

        dialog_buttons = QDialogButtonBox(QDialogButtonBox.Ok | QDialogButtonBox.Cancel)
        dialog_buttons.accepted.connect(dialog.accept)
        dialog_buttons.rejected.connect(dialog.reject)
        dialog_layout.addWidget(dialog_buttons)
        dialog.setLayout(dialog_layout)

        if dialog.exec() == QDialog.Accepted:
            self.current_analysis.start_date = date(date_edit_start.date().year(), date_edit_start.date().month(), date_edit_start.date().day())
            self.current_analysis.end_date = date(date_edit_end.date().year(), date_edit_end.date().month(), date_edit_end.date().day())
            self.plot_chart()

    def toggle_ma_period(self):
        period, ok = QInputDialog.getInt(self, "Período das Médias Móveis", "Período das Médias Móveis", self.current_settings.ma_period, 1, 1000, 1)
        if ok:
            self.current_settings.ma_period = period
            self.current_settings.ema_period = period
            self.current_settings.wma_period = period
            self.plot_chart()

    def toggle_candlestick_period(self):
        period, ok = QInputDialog.getInt(self, "Período dos Candlesticks", "Período (em dias):", value=self.current_settings.candlestick_period)
        if ok:
            self.current_settings.candlestick_period = period
            self.current_analysis.candlestick_period = period
            self.plot_chart()

    def set_ticker(self, ticker):
        self.current_analysis.ticker = ticker
        self.indicator_updater.update_indicators(ticker)
        self.update_menu_state()

    def update_menu_state(self):
        has_ticker = bool(self.current_analysis.ticker)
        self.edit_menu.setEnabled(has_ticker)
        self.view_menu.setEnabled(has_ticker)
        self.informacoes_menu.setEnabled(has_ticker)
        self.estrategias_menu.setEnabled(has_ticker)
        self.export_menu.setEnabled(has_ticker)

    def show_detailed_metrics(self):
        if self.current_analysis.ticker:
            self.metrics_window = MetricsWindow(self.current_analysis.ticker)
            self.metrics_window.show()
        else:
            QMessageBox.warning(self, "Erro", "Por favor, selecione um ticker primeiro.")

    def open_smart_metrics(self):
        if self.current_analysis.ticker:
            self.smart_metrics_window = SmartMetricsWindow(self.current_analysis.ticker)
            self.smart_metrics_window.show()
        else:
            QMessageBox.warning(self, "Erro", "Por favor, selecione um ticker primeiro.")

    def show_revenue_income_chart(self):
        if self.current_analysis.ticker:
            self.revenue_income_chart = RevenueIncomeChart(self.current_analysis.ticker)
            self.revenue_income_chart.show()
        else:
            QMessageBox.warning(self, "Erro", "Por favor, selecione um ticker primeiro.")

    def show_assets_liabilities_chart(self):
        if self.current_analysis.ticker:
            self.assets_liabilities_chart = AssetsLiabilitiesChart(self.current_analysis.ticker)
            self.assets_liabilities_chart.show()
        else:
            QMessageBox.warning(self, "Erro", "Por favor, selecione um ticker primeiro.")

    def export_to_excel(self):
        if not self.current_analysis.ticker:
            QMessageBox.warning(self, "Erro", "Por favor, selecione um ticker primeiro.")
            return

        try:
            data = DataFetcher.fetch_stock_data(
                self.current_analysis.ticker,
                self.current_analysis.start_date,
                self.current_analysis.end_date,
                self.current_analysis.candlestick_period,
                self.candlestick_cache
            )
            if data is None or data.empty:
                QMessageBox.warning(self, "Erro", "Não há dados para exportar.")
                return

            file_path, _ = QFileDialog.getSaveFileName(self, "Salvar dados como XLSX", "", "Excel Files (*.xlsx)")
            if file_path:
                data.to_excel(file_path, sheet_name=self.current_analysis.ticker)
                QMessageBox.information(self, "Sucesso", "Dados exportados com sucesso!")
        except Exception as e:
            QMessageBox.warning(self, "Erro", f"Erro ao exportar dados: {e}")

    def generate_intelligent_report_docx(self):
        if not self.current_analysis.ticker:
            QMessageBox.warning(self, "Erro", "Por favor, selecione um ticker primeiro.")
            return

        try:
            file_path, _ = QFileDialog.getSaveFileName(self, "Salvar relatório como DOCX", "", "Word Documents (*.docx)")
            if file_path:
                from docx import Document
                from docx.shared import Inches
                from docx.enum.text import WD_ALIGN_PARAGRAPH

                document = Document()
                document.add_heading(f"Relatório Inteligente - {self.current_analysis.ticker}", level=1)

                # Adicionar informações básicas
                document.add_heading("Informações Básicas", level=2)
                document.add_paragraph(f"Ticker: {self.current_analysis.ticker}")
                document.add_paragraph(f"Período de Análise: {self.current_analysis.start_date} - {self.current_analysis.end_date}")
                document.add_paragraph(f"Período dos Candlesticks: {self.current_analysis.candlestick_period} {'dia' if self.current_analysis.candlestick_period == 1 else 'dias'}")

                # Adicionar indicadores
                document.add_heading("Indicadores", level=2)
                document.add_paragraph(self.pvp_indicator.text())
                document.add_paragraph(self.pe_indicator.text())
                document.add_paragraph(self.roe_indicator.text())
                document.add_paragraph(self.dividend_yield_indicator.text())
                document.add_paragraph(self.debt_to_ebitda_indicator.text())
                document.add_paragraph(self.net_margin_indicator.text())

                # Adicionar gráfico de candlestick
                document.add_heading("Gráfico de Candlestick", level=2)
                temp_image_path = "temp_candlestick_chart.png"
                self.canvas.figure.savefig(temp_image_path, bbox_inches='tight')
                document.add_picture(temp_image_path, width=Inches(6))
                os.remove(temp_image_path)

                # Adicionar informações sobre médias móveis
                document.add_heading("Médias Móveis", level=2)
                if 'SMA' in self.current_analysis.medias:
                    document.add_paragraph(f"Média Móvel Simples (SMA): Ativada")
                if 'EMA' in self.current_analysis.medias:
                    document.add_paragraph(f"Média Móvel Exponencial (EMA): Ativada")
                if 'WMA' in self.current_analysis.medias:
                    document.add_paragraph(f"Média Móvel Ponderada (WMA): Ativada")

                # Adicionar informações sobre indicadores adicionais
                document.add_heading("Indicadores Adicionais", level=2)
                if self.current_analysis.show_volume:
                    document.add_paragraph("Volumes: Ativado")
                if self.current_analysis.show_ifr:
                    document.add_paragraph("Índice de Força Relativa (IFR): Ativado")
                if self.current_analysis.show_macd:
                    document.add_paragraph("Média Móvel de Convergência/Divergência (MACD): Ativado")
                if self.current_analysis.show_bandas_bollinger:
                    document.add_paragraph("Bandas de Bollinger: Ativadas")
                if self.current_analysis.show_estocastico_normal:
                    document.add_paragraph("Indicador Estocástico Normal: Ativado")
                if self.current_analysis.show_estocastico_lento:
                    document.add_paragraph("Indicador Estocástico Lento: Ativado")

                document.save(file_path)
                QMessageBox.information(self, "Sucesso", "Relatório gerado com sucesso!")
        except Exception as e:
            QMessageBox.warning(self, "Erro", f"Erro ao gerar relatório: {e}")

    def delayed_draw(self):
        if self.last_mouse_event:
            self.plot_chart()
            self.last_mouse_event = None

    def mouseMoveEvent(self, event):
        self.last_mouse_event = event
        self.mouse_move_timer.start(100)

    def desfazer(self):
        if self.current_analysis.plot_history:
            self.current_analysis.plot_history.pop()
            if self.current_analysis.plot_history:
                previous_state = self.current_analysis.plot_history[-1]
                self.current_analysis.medias = previous_state['medias']
                self.current_analysis.show_volume = previous_state['show_volume']
                self.current_analysis.show_ifr = previous_state['show_ifr']
                self.current_analysis.show_macd = previous_state['show_macd']
                self.current_analysis.show_bandas_bollinger = previous_state['show_bandas_bollinger']
                self.current_analysis.show_estocastico_normal = previous_state['show_estocastico_normal']
                self.current_analysis.show_estocastico_lento = previous_state['show_estocastico_lento']
                self.plot_chart()
            else:
                self.plot_chart()

    def closeEvent(self, event):
        self.current_settings.save_settings()
        event.accept()

    def toggle_sidebar(self):
        if not hasattr(self, 'dock_widget') or not self.dock_widget:
            self.dock_widget = QDockWidget("Principais Métricas", self)
            self.dock_widget.setAllowedAreas(Qt.LeftDockWidgetArea | Qt.RightDockWidgetArea)
            
            widget = QWidget()
            layout = QVBoxLayout(widget)
            layout.addWidget(self.pvp_indicator)
            layout.addWidget(self.pe_indicator)
            layout.addWidget(self.roe_indicator)
            layout.addWidget(self.dividend_yield_indicator)
            layout.addWidget(self.debt_to_ebitda_indicator)
            layout.addWidget(self.net_margin_indicator)
            
            self.dock_widget.setWidget(widget)
            self.addDockWidget(Qt.RightDockWidgetArea, self.dock_widget)
        else:
            if self.dock_widget.isVisible():
                self.dock_widget.hide()
            else:
                self.dock_widget.show()

if __name__ == '__main__':
    app = QApplication(sys.argv)
    mainWin = NovaGUI()
    mainWin.show()
    sys.exit(app.exec())